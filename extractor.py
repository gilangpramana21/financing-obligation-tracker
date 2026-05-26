"""LLM-powered extraction pipeline for financing agreement documents."""

import json
import os
from datetime import datetime, date
from pathlib import Path
from typing import Dict, List, Optional

import pdfplumber
try:
    import pymupdf
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
from docx import Document
from dotenv import load_dotenv

# Import LLM libraries
try:
    from google import genai
    from google.genai import types
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

from models import init_db, get_session, Agreement, ReportingObligation, Covenant, OtherObligation

# Load environment variables
load_dotenv()


class DocumentExtractor:
    """Extract structured obligation data from financing agreement documents."""
    
    def __init__(self, llm_provider: str = "gemini", api_key: Optional[str] = None):
        """
        Initialize extractor with specified LLM provider.
        
        Args:
            llm_provider: "gemini" (free, default) or "anthropic" (paid)
            api_key: API key for the provider (optional, reads from .env if not provided)
        """
        self.llm_provider = llm_provider.lower()
        
        if self.llm_provider == "gemini":
            if not GEMINI_AVAILABLE:
                raise ImportError("google-genai not installed. Install with: pip install google-genai")
            
            self.api_key = api_key or os.getenv('GEMINI_API_KEY')
            if not self.api_key:
                raise ValueError("GEMINI_API_KEY not found in environment variables")
            
            self.client = genai.Client(api_key=self.api_key)
            # Use latest stable models (confirmed available)
            self.model_name = "gemini-2.5-flash"  # Latest stable model
            self.fallback_model = "gemini-2.0-flash"  # Proven fallback
            print(f"🆓 Using {self.model_name} (Free) with fallback to {self.fallback_model}")
            
        elif self.llm_provider == "anthropic":
            if not ANTHROPIC_AVAILABLE:
                raise ImportError("anthropic not installed. Install with: pip install anthropic")
            
            self.api_key = api_key or os.getenv('ANTHROPIC_API_KEY')
            if not self.api_key:
                raise ValueError("ANTHROPIC_API_KEY not found in environment variables")
            
            self.client = anthropic.Anthropic(api_key=self.api_key)
            self.model_name = "claude-sonnet-4"
            self.fallback_model = None
            print("💰 Using Claude Sonnet 4 (Paid)")
            
        else:
            raise ValueError(f"Unsupported LLM provider: {llm_provider}. Use 'gemini' or 'anthropic'")
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using pdfplumber."""
        text_parts = []
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            return "\n\n".join(text_parts)
        
        except Exception as e:
            print(f"⚠️  pdfplumber failed, trying PyPDF2: {e}")
            # Fallback to PyPDF2
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(pdf_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return "\n\n".join(text_parts)
            except Exception as e2:
                # Last resort: try PyMuPDF if available
                if PYMUPDF_AVAILABLE:
                    try:
                        doc = pymupdf.open(pdf_path)
                        for page in doc:
                            text_parts.append(page.get_text())
                        doc.close()
                        return "\n\n".join(text_parts)
                    except Exception as e3:
                        raise Exception(f"Failed to extract text from PDF: {e3}")
                else:
                    raise Exception(f"Failed to extract text from PDF: {e2}")
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from Word document."""
        doc = Document(docx_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from document (PDF or Word)."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif suffix in ['.docx', '.doc']:
            return self.extract_text_from_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def build_extraction_prompt(self, document_text: str) -> str:
        """Build prompt for LLM to extract structured obligation data."""
        
        prompt = f"""You are analyzing a financing agreement document. Extract all obligations, covenants, and reporting requirements into a structured JSON format.

DOCUMENT TEXT:
{document_text}

---

Extract the following information and return ONLY valid JSON (no markdown, no explanation):

{{
  "financier": "string (name of the lender/financier)",
  "agreement_name": "string (full name of the agreement)",
  "contract_start": "YYYY-MM-DD (contract start date)",
  "contract_end": "YYYY-MM-DD (contract end/maturity date)",
  "facility_amount": number (loan/facility amount as a number),
  "currency": "string (IDR, USD, EUR, etc.)",
  "reporting_obligations": [
    {{
      "report_name": "string (name of the report)",
      "frequency": "monthly|quarterly|semi-annual|annual",
      "due_day": number (days after period end when report is due),
      "description": "string (detailed description of reporting requirement)",
      "next_due": "YYYY-MM-DD (calculate next due date based on frequency)"
    }}
  ],
  "covenants": [
    {{
      "name": "string (covenant name)",
      "type": "minimum|maximum",
      "metric": "string (what is being measured)",
      "threshold": number (the threshold value),
      "unit": "IDR|USD|ratio|percent",
      "description": "string (full covenant description)"
    }}
  ],
  "other_obligations": [
    {{
      "category": "Notification|Approval Required|Restriction|Action Required",
      "description": "string (detailed description)",
      "is_ongoing": true
    }}
  ]
}}

EXTRACTION RULES:
1. For dates: Use ISO format YYYY-MM-DD. If exact date is unclear, estimate based on context.
2. For amounts: Extract as pure numbers (e.g., 50000000000 not "50 billion")
3. For frequencies: Use exact values: "monthly", "quarterly", "semi-annual", or "annual"
4. For covenant types: "minimum" for floor covenants, "maximum" for ceiling covenants
5. For other_obligations categories: Choose from: "Notification", "Approval Required", "Restriction", "Action Required"
6. If information is missing or unclear, use reasonable defaults or null
7. Extract ALL obligations found in the document - be comprehensive
8. For next_due dates: Calculate based on current date ({date.today().isoformat()}) and frequency

Return ONLY the JSON object, no other text."""

        return prompt
    
    def extract_obligations_from_text(self, document_text: str, max_retries: int = 3) -> Dict:
        """Use LLM to extract structured obligation data from document text with retry logic."""
        
        import time
        
        prompt = self.build_extraction_prompt(document_text)
        
        # Try primary model first, then fallback model if available
        models_to_try = [self.model_name]
        if hasattr(self, 'fallback_model') and self.fallback_model:
            models_to_try.append(self.fallback_model)
        
        for model_name in models_to_try:
            last_error = None
            
            if model_name != self.model_name:
                print(f"   🔄 Trying fallback model: {model_name}")
            
            for attempt in range(max_retries):
                try:
                    if self.llm_provider == "gemini":
                        # Gemini API call (new SDK)
                        response = self.client.models.generate_content(
                            model=model_name,
                            contents=prompt
                        )
                        response_text = response.text.strip()
                        
                    elif self.llm_provider == "anthropic":
                        # Anthropic API call
                        message = self.client.messages.create(
                            model=model_name,
                            max_tokens=8000,
                            temperature=0,
                            messages=[
                                {
                                    "role": "user",
                                    "content": prompt
                                }
                            ]
                        )
                        response_text = message.content[0].text.strip()
                    
                    # Remove markdown code blocks if present
                    if response_text.startswith('```'):
                        lines = response_text.split('\n')
                        # Remove first line (```json or ```) and last line (```)
                        response_text = '\n'.join(lines[1:-1]) if len(lines) > 2 else response_text
                    
                    # Parse JSON
                    extracted_data = json.loads(response_text)
                    
                    # Success! Return the data
                    if attempt > 0 or model_name != self.model_name:
                        print(f"   ✓ Succeeded with {model_name} on attempt {attempt + 1}")
                    return extracted_data
                
                except json.JSONDecodeError as e:
                    # JSON parsing error - don't retry, this is a different issue
                    raise Exception(f"Failed to parse LLM response as JSON: {e}\nResponse: {response_text[:500]}...")
                
                except Exception as e:
                    last_error = e
                    error_str = str(e)
                    
                    # Check if it's a rate limit or availability error
                    if '503' in error_str or '429' in error_str or 'UNAVAILABLE' in error_str or 'RESOURCE_EXHAUSTED' in error_str:
                        if attempt < max_retries - 1:
                            # Exponential backoff: 2^attempt seconds (2s, 4s, 8s)
                            wait_time = 2 ** (attempt + 1)
                            print(f"   ⚠️  {model_name} overloaded (attempt {attempt + 1}/{max_retries}). Retrying in {wait_time}s...")
                            time.sleep(wait_time)
                            continue
                        else:
                            # This model exhausted all retries, try next model
                            print(f"   ❌ {model_name} unavailable after {max_retries} attempts")
                            break
                    else:
                        # Other error - don't retry
                        raise Exception(f"LLM extraction failed: {e}")
            
            # If we get here, this model failed all retries, try next model
            if model_name == models_to_try[-1]:
                # This was the last model
                raise Exception(f"All models failed. Last error: {last_error}")
        
        # If we get here, all models and retries failed
        raise Exception(f"LLM extraction failed after trying all available models")
    
    def validate_extracted_data(self, data: Dict) -> bool:
        """Validate extracted data structure."""
        required_fields = ['financier', 'agreement_name', 'contract_start', 
                          'contract_end', 'facility_amount', 'currency']
        
        for field in required_fields:
            if field not in data or data[field] is None:
                raise ValueError(f"Missing required field: {field}")
        
        # Validate date formats
        for date_field in ['contract_start', 'contract_end']:
            try:
                datetime.strptime(data[date_field], '%Y-%m-%d')
            except ValueError:
                raise ValueError(f"Invalid date format for {date_field}: {data[date_field]}")
        
        # Validate lists exist
        if 'reporting_obligations' not in data:
            data['reporting_obligations'] = []
        if 'covenants' not in data:
            data['covenants'] = []
        if 'other_obligations' not in data:
            data['other_obligations'] = []
        
        return True
    
    def save_to_database(self, extracted_data: Dict, db_path: str = 'obligation_tracker.db'):
        """Save extracted obligation data to database."""
        
        # Validate data
        self.validate_extracted_data(extracted_data)
        
        # Initialize database
        engine = init_db(db_path)
        session = get_session(engine)
        
        try:
            # Create agreement
            agreement = Agreement(
                financier=extracted_data['financier'],
                agreement_name=extracted_data['agreement_name'],
                contract_start=datetime.strptime(extracted_data['contract_start'], '%Y-%m-%d').date(),
                contract_end=datetime.strptime(extracted_data['contract_end'], '%Y-%m-%d').date(),
                facility_amount=float(extracted_data['facility_amount']),
                currency=extracted_data['currency']
            )
            session.add(agreement)
            session.flush()
            
            # Add reporting obligations
            for report in extracted_data['reporting_obligations']:
                reporting = ReportingObligation(
                    agreement_id=agreement.id,
                    report_name=report['report_name'],
                    frequency=report['frequency'],
                    due_day=report['due_day'],
                    description=report.get('description', ''),
                    next_due=datetime.strptime(report['next_due'], '%Y-%m-%d').date()
                )
                session.add(reporting)
            
            # Add covenants
            for covenant in extracted_data['covenants']:
                cov = Covenant(
                    agreement_id=agreement.id,
                    name=covenant['name'],
                    type=covenant['type'],
                    metric=covenant['metric'],
                    threshold=float(covenant['threshold']),
                    unit=covenant['unit'],
                    description=covenant.get('description', ''),
                    current_value=None,  # To be updated separately
                    last_updated=None
                )
                session.add(cov)
            
            # Add other obligations
            for obligation in extracted_data['other_obligations']:
                other = OtherObligation(
                    agreement_id=agreement.id,
                    category=obligation['category'],
                    description=obligation['description'],
                    is_ongoing=obligation.get('is_ongoing', True)
                )
                session.add(other)
            
            session.commit()
            
            print(f"✅ Successfully saved agreement: {agreement.agreement_name}")
            print(f"   📊 {len(extracted_data['reporting_obligations'])} reporting obligations")
            print(f"   📈 {len(extracted_data['covenants'])} covenants")
            print(f"   📋 {len(extracted_data['other_obligations'])} other obligations")
            
            return agreement.id
        
        except Exception as e:
            session.rollback()
            raise Exception(f"Failed to save to database: {e}")
        finally:
            session.close()
    
    def process_document(self, file_path: str, db_path: str = 'obligation_tracker.db') -> int:
        """
        Complete pipeline: extract text → LLM extraction → save to database.
        
        Returns: agreement_id
        """
        print(f"📄 Processing document: {file_path}")
        
        # Step 1: Extract text
        print("   Extracting text from document...")
        document_text = self.extract_text(file_path)
        print(f"   ✓ Extracted {len(document_text)} characters")
        
        # Step 2: LLM extraction
        print("   Analyzing with Claude...")
        extracted_data = self.extract_obligations_from_text(document_text)
        print(f"   ✓ Extracted obligations from {extracted_data['financier']}")
        
        # Step 3: Save to database
        print("   Saving to database...")
        agreement_id = self.save_to_database(extracted_data, db_path)
        
        print(f"✅ Document processing complete (Agreement ID: {agreement_id})")
        return agreement_id


def process_documents_folder(folder_path: str = 'documents', db_path: str = 'obligation_tracker.db', llm_provider: str = "gemini"):
    """Process all PDF and Word documents in a folder."""
    
    folder = Path(folder_path)
    if not folder.exists():
        print(f"Creating documents folder: {folder_path}")
        folder.mkdir(parents=True, exist_ok=True)
        return
    
    # Find all documents
    documents = list(folder.glob('*.pdf')) + list(folder.glob('*.docx')) + list(folder.glob('*.doc'))
    
    if not documents:
        print(f"No documents found in {folder_path}")
        return
    
    print(f"Found {len(documents)} document(s) to process\n")
    
    extractor = DocumentExtractor(llm_provider=llm_provider)
    
    for doc_path in documents:
        try:
            extractor.process_document(str(doc_path), db_path)
            print()
        except Exception as e:
            print(f"❌ Failed to process {doc_path.name}: {e}\n")


if __name__ == '__main__':
    import sys
    
    # Default to Gemini (free)
    llm_provider = "gemini"
    file_path = None
    
    # Parse command line arguments
    if len(sys.argv) > 1:
        for arg in sys.argv[1:]:
            if arg.lower() in ['--anthropic', '--claude']:
                llm_provider = "anthropic"
            elif arg.lower() in ['--gemini', '--google']:
                llm_provider = "gemini"
            elif not arg.startswith('--'):
                file_path = arg
    
    if file_path:
        # Process specific file
        extractor = DocumentExtractor(llm_provider=llm_provider)
        extractor.process_document(file_path)
    else:
        # Process all documents in folder
        process_documents_folder(llm_provider=llm_provider)
